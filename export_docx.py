from __future__ import annotations

from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def _safe_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, list):
        return "\n".join([str(v) for v in value])
    if isinstance(value, dict):
        return "\n".join([f"{k}: {v}" for k, v in value.items()])
    return str(value)


def _add_table(doc: Document, rows: list[dict], columns: list[str]):
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for i, col in enumerate(columns):
        cell = table.rows[0].cells[i]
        cell.text = col
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for row in rows or []:
        cells = table.add_row().cells
        for i, col in enumerate(columns):
            cells[i].text = _safe_text(row.get(col, "")) if isinstance(row, dict) else _safe_text(row)
    doc.add_paragraph("")
    return table


def _selected_tools_rows(state: dict) -> list[dict]:
    rec = state.get("tool_recommendation", {}) or {}
    selected = rec.get("selected_tools", []) or []
    candidates = rec.get("candidates", []) or []
    rows = []
    for sel in selected:
        if isinstance(sel, dict):
            rows.append(sel)
            continue
        match = next((c for c in candidates if isinstance(c, dict) and c.get("name") == sel), None)
        if match:
            rows.append(match)
        else:
            rows.append({"name": str(sel), "reason": ""})
    return rows


def _selected_model_rows(state: dict) -> list[dict]:
    model = (state.get("model_recommendation", {}) or {}).get("selected_model", {}) or {}
    if not isinstance(model, dict) or not model:
        return []
    return [{"항목": k, "내용": v} for k, v in model.items()]


def _worksheet_rows(questions: list[Any]) -> list[dict]:
    # 화면 렌더러와 완전히 같지는 않지만, 문제/보기/답칸을 분리하려고 노력함.
    try:
        from worksheet_renderer import normalize_question_object, ensure_list
    except Exception:
        normalize_question_object = None
        ensure_list = lambda x: x if isinstance(x, list) else ([x] if x else [])
    rows = []
    for i, q in enumerate(questions or [], start=1):
        if normalize_question_object:
            item = normalize_question_object(q, i)
            body = item.get("question_text", "")
            support = "\n".join(ensure_list(item.get("supporting_lines", [])))
            choices = "\n".join(ensure_list(item.get("choices", [])))
            rows.append({"번호": i, "문제": body, "보기/자료": support or choices, "답칸": "(      )"})
        else:
            rows.append({"번호": i, "문제": _safe_text(q), "보기/자료": "", "답칸": "(      )"})
    return rows


def create_full_docx(state: dict) -> BytesIO:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Malgun Gothic"
    style.font.size = Pt(10)

    title = doc.add_heading("AI 기반 수업 설계 결과", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    lesson = state.get("lesson_input", {}) or {}
    teacher_intent = state.get("teacher_intent", {}) or {}

    doc.add_heading("1. 기본 정보", 1)
    _add_table(doc, [
        {"항목": "과목", "내용": lesson.get("subject", "")},
        {"항목": "학년", "내용": lesson.get("grade", "")},
        {"항목": "학기", "내용": lesson.get("semester", "")},
        {"항목": "단원", "내용": lesson.get("unit", "")},
        {"항목": "차시 주제", "내용": lesson.get("lesson_topic", "")},
        {"항목": "학습 목표", "내용": lesson.get("learning_objectives", "")},
        {"항목": "성취기준", "내용": lesson.get("curriculum_text", "")},
        {"항목": "수업자 의도", "내용": teacher_intent.get("raw_text", "")},
    ], ["항목", "내용"])

    doc.add_heading("2. 선택된 교구", 1)
    tool_rows = _selected_tools_rows(state)
    if tool_rows:
        _add_table(doc, tool_rows, ["name", "score_total", "reason", "best_for", "cautions"])
    else:
        doc.add_paragraph("선택된 교구 없음")

    doc.add_heading("3. 선택된 학습모형", 1)
    model_rows = _selected_model_rows(state)
    if model_rows:
        _add_table(doc, model_rows, ["항목", "내용"])
    else:
        doc.add_paragraph("선택된 학습모형 없음")

    lesson_plan = state.get("lesson_plan_output", {}) or {}
    doc.add_heading("4. 지도안 개요", 1)
    overview = lesson_plan.get("overview", {}) or {}
    _add_table(doc, [{"항목": k, "내용": v} for k, v in overview.items()], ["항목", "내용"])

    doc.add_heading("5. 표 형식 지도안", 1)
    _add_table(doc, lesson_plan.get("lesson_table", []) or [], ["단계", "시량", "교사 활동", "학생 활동", "자료 및 유의점"])

    doc.add_heading("6. 평가 계획", 1)
    assess = lesson_plan.get("assessment_plan", []) or []
    if assess:
        # 키가 상/중/하 또는 평가기준 중 무엇이든 가능한 한 보존
        cols = ["평가시기", "평가내용", "평가방법", "상", "중", "하"]
        if not any(c in assess[0] for c in ["상", "중", "하"]):
            cols = ["평가시기", "평가내용", "평가방법", "평가기준"]
        _add_table(doc, assess, cols)

    doc.add_heading("7. 피드백 전략", 1)
    _add_table(doc, lesson_plan.get("feedback_strategy", []) or [], ["성취수준", "평가기준", "피드백 전략"])

    slides = (state.get("slides_output", {}) or {}).get("slides", []) or []
    doc.add_heading("8. 슬라이드", 1)
    if slides:
        # 시각화 프롬프트 보존
        normalized = []
        for s in slides:
            if isinstance(s, dict):
                row = dict(s)
                row["image_prompt"] = row.get("image_prompt") or row.get("visuals", "")
                normalized.append(row)
        _add_table(doc, normalized, ["slide_no", "title", "purpose", "main_content", "teacher_prompt", "image_prompt"])

    worksheet = state.get("worksheet_output", {}) or {}
    doc.add_heading("9. 학습지", 1)
    for level_key, label in [("high", "상"), ("middle", "중"), ("low", "하")]:
        data = worksheet.get(level_key, {}) if isinstance(worksheet, dict) else {}
        questions = data.get("questions", []) if isinstance(data, dict) else (data if isinstance(data, list) else [])
        if questions:
            doc.add_heading(label, 2)
            _add_table(doc, _worksheet_rows(questions), ["번호", "문제", "보기/자료", "답칸"])

    answers = worksheet.get("answers", []) if isinstance(worksheet, dict) else []
    if answers:
        doc.add_heading("정답/예시답안", 2)
        for a in answers:
            doc.add_paragraph(_safe_text(a))

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
